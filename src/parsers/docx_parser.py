"""DOCX parser — paragraph batches + tables via python-docx."""
import logging
from typing import Generator

from src.parsers.base_parser import BaseParser
from src.exceptions import ParseError

log = logging.getLogger(__name__)
_PARAGRAPH_BATCH = 20  # paragraphs yielded per segment


class DocxParser(BaseParser):
    """
    Reads paragraphs in batches of _PARAGRAPH_BATCH and also extracts
    table cell content as pipe-separated rows. Falls back to full-body
    XML text extraction for DOCX files whose content lives outside normal
    paragraphs (text boxes, headers embedded in body, etc.).
    """

    def supports(self, extension: str) -> bool:
        return extension.lower() == ".docx"

    def parse(self, file_path: str) -> Generator[str, None, None]:
        try:
            from docx import Document
            from docx.oxml.ns import qn
        except ImportError:
            raise ParseError("python-docx is not installed. Run: pip install python-docx")

        try:
            doc = Document(file_path)
        except Exception as e:
            raise ParseError(f"Failed to open DOCX '{file_path}': {e}") from e

        total_paras = len(doc.paragraphs)
        non_empty_paras = sum(1 for p in doc.paragraphs if p.text.strip())
        total_tables = len(doc.tables)
        log.info(f"[DocxParser] Total paragraphs={total_paras} | Non-empty={non_empty_paras} | Tables={total_tables}")

        # ── Paragraphs ────────────────────────────────────────────────────
        batch: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                batch.append(text)
            if len(batch) >= _PARAGRAPH_BATCH:
                yield "\n".join(batch)
                batch = []
        if batch:
            yield "\n".join(batch)

        # ── Tables ────────────────────────────────────────────────────────
        table_rows_yielded = 0
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    yield " | ".join(cells)
                    table_rows_yielded += 1
        log.info(f"[DocxParser] Table rows yielded={table_rows_yielded}")

        # ── Fallback: text boxes and shapes ───────────────────────────────
        # Some DOCX files store content in drawing/textbox elements that
        # are invisible to doc.paragraphs. Extract all <w:t> text nodes
        # from the body XML directly.
        extra_texts: list[str] = []
        try:
            for elem in doc.element.body.iter(qn("w:t")):
                t = (elem.text or "").strip()
                if t:
                    extra_texts.append(t)
        except Exception as e:
            log.warning(f"[DocxParser] XML fallback failed: {e}")

        # Only yield extras that weren't already covered by paragraphs/tables
        para_text_set = set()
        for p in doc.paragraphs:
            para_text_set.update(p.text.split())

        extra_filtered = [t for t in extra_texts if t not in para_text_set]
        log.info(f"[DocxParser] Extra text nodes from XML fallback: {len(extra_filtered)}")

        if extra_filtered:
            # Batch them up
            batch = []
            for t in extra_filtered:
                batch.append(t)
                if len(batch) >= _PARAGRAPH_BATCH:
                    yield "\n".join(batch)
                    batch = []
            if batch:
                yield "\n".join(batch)
